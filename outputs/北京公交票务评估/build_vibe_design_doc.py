from __future__ import annotations

from collections import defaultdict
from datetime import date
from pathlib import Path
from re import sub
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


SOURCE = Path("/Users/mate/Docs/Work/北京公交集团票务综合管理平台系统升级-需求分析说明书（初稿0506）.docx")
OUT = Path("/Users/mate/Codes/mate/mateclaw/outputs/北京公交票务评估/北京公交集团票务综合管理平台系统升级-VibeCoding实施方案与工作量清单.docx")

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(11, 37, 69)
MUTED = RGBColor(96, 108, 122)
LIGHT = "F2F4F7"
LIGHT_BLUE = "E8EEF5"
CALLOUT = "F4F6F9"


def clean(text: str) -> str:
    text = text.replace("\u200c", "").replace("\u200b", "").replace("\ufeff", "")
    text = sub(r"\s+", " ", text).strip()
    return text


def module_name(text: str) -> str:
    text = clean(text)
    text = sub(r"^[\d.、\s]+", "", text)
    return text.strip("‌ ")


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def shade(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for key, val in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")


def table_width(table, widths: list[float]):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = Inches(widths[idx])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            margins(cell)


def add_para(doc, text="", bold=False, color=None, size=10.5, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.12
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def heading(doc, text, level):
    p = doc.add_heading(text, level=level)
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(8 if level == 1 else 5)
    for r in p.runs:
        set_run_font(r, size={1: 16, 2: 13, 3: 12}.get(level, 11), bold=True, color=BLUE if level < 3 else DARK_BLUE)


def bullets(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        set_run_font(r, size=10)


def numbered(doc, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.12
        r = p.add_run(item)
        set_run_font(r, size=10)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[float], fill=LIGHT, font_size=8.8):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        c = table.rows[0].cells[i]
        c.text = h
        shade(c, fill)
        margins(c)
        for p in c.paragraphs:
            p.paragraph_format.space_after = Pt(0)
            for r in p.runs:
                set_run_font(r, size=font_size, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = str(value)
            margins(cells[i])
            for p in cells[i].paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.02
                for r in p.runs:
                    set_run_font(r, size=font_size)
    table_width(table, widths)
    doc.add_paragraph()


def callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade(cell, CALLOUT)
    margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(body)
    set_run_font(r2, size=10)
    table_width(table, [6.5])
    doc.add_paragraph()


def configure(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.45)
    section.footer_distance = Inches(0.45)
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12
    for style_name, size, color in [("Heading 1", 16, BLUE), ("Heading 2", 13, BLUE), ("Heading 3", 12, DARK_BLUE)]:
        st = styles[style_name]
        st.font.name = "Calibri"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = color
    header = section.header.paragraphs[0]
    header.text = "北京公交集团票务综合管理平台系统升级 | Vibe Coding 实施与工作量评估"
    for r in header.runs:
        set_run_font(r, size=8.5, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "内部评估文件"
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for r in footer.runs:
        set_run_font(r, size=8.5, color=MUTED)


def extract_items():
    doc = Document(SOURCE)
    table = doc.tables[38]
    items = []
    current = ""
    for row in table.rows[1:]:
        cells = [clean(c.text) for c in row.cells[:5]]
        mod, page, func, change, note = cells
        if mod:
            current = module_name(mod)
        if not page:
            continue
        items.append({
            "module": current,
            "page": page,
            "function": func,
            "change": change,
            "note": note,
        })
    return items


def classify(item):
    text = f"{item['module']} {item['page']} {item['function']} {item['change']} {item['note']}"
    if "可删除" in text or "删除" in item["page"] or item["change"] == "删除":
        return "拟删除/整合", 0, 1, "确认下线、菜单隐藏、历史数据保留和跳转清理。"
    if any(k in text for k in ["结算", "重新结算", "配票", "调票", "退票", "库存", "入库", "出库", "核销", "调拨", "印制", "票号"]):
        lo, hi = 8, 18
        if any(k in text for k in ["票单结算", "重新结算", "票袋调票", "票柜退票", "票袋退票"]):
            lo, hi = 14, 26
        return "核心业务", lo, hi, "需要状态机、库存流水、票号段校验、事务一致性和回归测试。"
    if any(k in item["module"] for k in ["统计报表"]) or "报表" in text or "统计" in text:
        return "报表统计", 6, 14, "需确认统计口径、权限范围、导出格式和历史数据性能。"
    if "导入" in text or "上传" in text or "下载" in text:
        return "导入导出", 5, 10, "需模板、字段校验、失败明细、导入日志和重复处理。"
    if any(k in item["module"] for k in ["系统管理"]) or any(k in item["page"] for k in ["用户", "权限", "菜单", "日志", "阈值"]):
        return "平台能力", 6, 14, "需统一认证授权、按钮权限、数据范围和审计记录。"
    if any(k in text for k in ["查询、新建、修改、删除", "添加、编辑、删除", "添加、编辑", "新增", "管理"]):
        return "常规功能", 4, 8, "以列表、表单、校验、权限按钮和操作日志为主。"
    return "查询维护", 3, 6, "以查询、筛选、详情、导出或基础维护为主。"


def build_rows(items):
    rows = []
    for idx, item in enumerate(items, start=1):
        kind, lo, hi, detail = classify(item)
        vibe_lo = round(lo * 0.65, 1)
        vibe_hi = round(hi * 0.75, 1)
        dev = build_dev_scope(item, kind)
        rows.append({
            **item,
            "idx": idx,
            "kind": kind,
            "lo": lo,
            "hi": hi,
            "vibe_lo": vibe_lo,
            "vibe_hi": vibe_hi,
            "dev": dev,
            "detail": detail,
        })
    return rows


def build_dev_scope(item, kind):
    page = item["page"]
    if kind == "拟删除/整合":
        return f"确认 {page} 下线或合并路径，处理菜单、权限、历史查询入口和数据保留。"
    if kind == "核心业务":
        return f"建设 {page} 的单据、审批/确认、库存流水、票号段校验、查询导出和异常回滚。"
    if kind == "报表统计":
        return f"建设 {page} 查询、统计聚合、权限过滤、Excel 导出、必要趋势展示和口径校验。"
    if kind == "导入导出":
        return f"建设 {page} 模板下载、文件上传、字段校验、入库处理、失败明细和导入日志。"
    if kind == "平台能力":
        return f"建设 {page} 的配置管理、权限控制、审计记录、查询筛选和维护页面。"
    return f"建设 {page} 的列表查询、详情、新增/编辑/停用或删除、权限按钮、日志和导出能力。"


def module_summary(rows):
    result = []
    grouped = defaultdict(list)
    for row in rows:
        grouped[row["module"]].append(row)
    for mod, vals in grouped.items():
        lo = sum(v["lo"] for v in vals)
        hi = sum(v["hi"] for v in vals)
        vlo = sum(v["vibe_lo"] for v in vals)
        vhi = sum(v["vibe_hi"] for v in vals)
        result.append([mod, str(len(vals)), f"{lo}-{hi}", f"{round(vlo)}-{round(vhi)}", module_focus(mod)])
    return result


def module_focus(mod):
    if "票单" in mod:
        return "票袋、配票、结算、重新结算、票号追溯，是核心复杂域。"
    if "票柜" in mod or "票库" in mod or "申请" in mod or "审核" in mod:
        return "围绕库存流转、审批状态和票号段控制。"
    if "统计报表" in mod:
        return "以口径确认、聚合性能、导出格式和权限过滤为重点。"
    if "系统" in mod:
        return "全局权限、菜单、日志、阈值，是平台底座。"
    if "基础" in mod:
        return "主数据维护和外部同步，是后续业务依赖。"
    return "按业务场景完成查询、维护、导入导出和日志审计。"


def make_landscape(doc):
    section = doc.add_section(WD_SECTION.NEW_PAGE)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.55)
    section.right_margin = Inches(0.55)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    return section


def build():
    items = extract_items()
    rows = build_rows(items)
    normal_lo = sum(r["lo"] for r in rows)
    normal_hi = sum(r["hi"] for r in rows)
    vibe_lo = round(sum(r["vibe_lo"] for r in rows))
    vibe_hi = round(sum(r["vibe_hi"] for r in rows))

    doc = Document()
    configure(doc)

    add_para(doc, "项目实施方案", bold=True, color=MUTED, size=11)
    p = doc.add_paragraph()
    r = p.add_run("北京公交集团票务综合管理平台系统升级")
    set_run_font(r, size=23, bold=True, color=INK)
    p.paragraph_format.space_after = Pt(4)
    p = doc.add_paragraph()
    r = p.add_run("Vibe Coding 实施方案与工作量清单")
    set_run_font(r, size=15, color=MUTED)
    p.paragraph_format.space_after = Pt(14)
    add_table(doc, ["字段", "内容"], [
        ["文档用途", "用于项目立项、排期、报价拆分、研发组织和验收范围确认"],
        ["编制日期", date.today().isoformat()],
        ["需求来源", "北京公交集团票务综合管理平台系统升级需求分析说明书（初稿0506）"],
        ["估算口径", "包含需求深化、设计、开发、自测、联调、迁移、测试配合和上线支持"],
    ], [1.35, 5.15], fill=LIGHT_BLUE, font_size=9)

    callout(doc, "简要结论", f"本项目按完整重构口径约 {len(rows)} 个功能项，正常开发工作量约 {normal_lo}-{normal_hi} 人日。采用 Vibe Coding 协同方式后，可将重复代码、测试样例、文档和脚手架类工作压缩，执行工作量约 {vibe_lo}-{vibe_hi} 人日，但票据规则、报表口径、数据迁移和上线验收仍需人工主导。")

    heading(doc, "一、简要评估", 1)
    add_table(doc, ["项目", "建议口径"], [
        ["功能规模", f"{len(rows)} 个功能项，覆盖 14 个一级模块，含基础资料、票库、票柜、票单、退票、充值、报表、系统管理等。"],
        ["正常团队", "10-12 人，8-10 个月完成；包含产品、前端、后端、数据迁移、测试、实施。"],
        ["Vibe Coding 团队", "7-9 人，7-8 个月完成；前提是需求冻结、接口材料及时、测试自动化和人工评审到位。"],
        ["核心难点", "票号段连续性、库存流水、票单结算、报表口径、数据权限、历史数据迁移、国产化适配。"],
        ["报价建议", "完整建设建议按 450-650 万区间，稳妥报价点 520-580 万；电子签章、等保测评、第三方授权另列。"],
    ], [1.4, 5.1], fill=LIGHT_BLUE, font_size=9)

    heading(doc, "二、模块工作量简表", 1)
    add_table(doc, ["模块", "功能项数", "正常人日", "协同后人日", "重点说明"], module_summary(rows), [1.25, 0.65, 0.85, 0.9, 2.85], fill=LIGHT)

    heading(doc, "三、Vibe Coding 实施方式", 1)
    add_para(doc, "本方案中的 Vibe Coding 不是直接让工具一次性生成完整系统，而是把需求、设计、编码、测试、文档拆成小批量可验证任务。每个功能都按照“需求卡片 - 数据结构 - 接口 - 页面 - 测试 - 评审 - 合并”的节奏推进。")
    numbered(doc, [
        "需求卡片化：每个页面或业务动作形成独立任务，明确角色、数据范围、输入输出、状态和验收标准。",
        "先底座后业务：先完成认证、菜单、权限、字典、日志、导入导出、文件存储等平台能力。",
        "核心域手工建模：票据库存、票号段、结算、退票、调账等领域模型由架构师和业务人员先确认，再进入编码。",
        "批量生成常规代码：常规列表、表单、接口、DTO、Mapper、权限按钮、导入模板可以批量生成并统一 review。",
        "测试前置：每个核心业务先写状态流转和金额/张数计算测试，避免只实现页面不验证账务逻辑。",
        "小步交付：每个模块独立演示、独立验收、独立回归，不把问题积压到总体验收阶段。",
    ])

    heading(doc, "四、实施步骤简表", 1)
    add_table(doc, ["阶段", "周期", "主要工作", "交付物"], [
        ["0. 需求深化", "3-5 周", "需求清单、业务流程、角色权限、报表口径、接口范围、迁移对象确认。", "需求矩阵、原型、报表口径表、接口清单、迁移清单。"],
        ["1. 平台底座", "4-6 周", "登录、用户、角色、菜单、数据权限、日志、字典、阈值、消息、文件导入导出。", "可登录可授权的基础系统。"],
        ["2. 基础与库存", "8-10 周", "基础资料、票库、票柜、审批、库存余额和流水。", "库存闭环 MVP。"],
        ["3. 票单核心", "8-12 周", "票袋、配票、结算、重新结算、退票、票号查询。", "票务室核心业务闭环。"],
        ["4. 扩展业务与报表", "8-10 周", "无人售、充值、特殊业务、集团/分公司/票务室报表。", "完整业务与统计能力。"],
        ["5. 联调上线", "8-10 周", "数据迁移、接口联调、UAT、性能、安全整改、上线演练。", "生产上线版本与运维交接材料。"],
    ], [0.9, 0.8, 3.05, 1.75], fill=LIGHT)

    heading(doc, "五、关键设计细节", 1)
    heading(doc, "5.1 库存与票号设计", 2)
    bullets(doc, [
        "采用“库存余额表 + 库存流水表 + 业务单据表”三层模型，不以单一库存字段承载全部业务历史。",
        "票号段字段统一包含票价、票种、票组、起号、止号、张数、状态、库存主体、来源单据。",
        "所有出入库动作必须写流水，流水记录来源、目标、数量、金额、操作人、操作时间和业务原因。",
        "连续票号段需校验重叠、断号、越界、已核销、已退票、已结算等状态。",
        "重新结算不得覆盖原记录，应形成更正单据和差异流水，保证历史可追溯。",
    ])
    heading(doc, "5.2 权限与数据范围", 2)
    bullets(doc, [
        "权限拆分为菜单权限、按钮权限、接口权限、数据范围权限和导出权限。",
        "集团管理员可分配全局权限；集团操作员按授权查看部分分公司；分公司角色默认本分公司；票务室角色只看本票务室。",
        "后端查询必须统一注入数据范围条件，不能只依赖前端隐藏菜单。",
        "导出、打印、报表接口必须复用同一套数据权限规则。",
    ])
    heading(doc, "5.3 报表设计", 2)
    bullets(doc, [
        "每张报表上线前先确认指标定义、筛选条件、统计维度、取数来源、金额精度、导出格式。",
        "集团和分公司同类报表尽量合并，用数据权限和维度字段控制展示结果。",
        "高频报表建立日汇总/月汇总表，避免每次查询实时扫大量流水。",
        "报表验收必须使用旧系统样例数据做金额和张数比对。",
    ])
    heading(doc, "5.4 数据迁移", 2)
    bullets(doc, [
        "迁移对象包括 Oracle 表、序列、视图、函数、历史业务数据和必要附件。",
        "先做数据体检，输出重复票号、缺失组织、异常状态、金额不平、非法日期等问题清单。",
        "至少执行三轮迁移演练：开发、测试、准生产，每轮生成记录数、金额、票号覆盖和异常清单。",
        "上线割接需要冻结旧系统写入、执行增量迁移、业务抽样核验，并保留回滚方案。",
    ])

    heading(doc, "六、部署与可运维性设计", 1)
    add_para(doc, "本项目属于票务核心管理系统，部署和运维能力应在一期同步设计，不能等到开发完成后补。运维设计目标是：可部署、可监控、可审计、可备份、可恢复、可灰度、可回滚。")
    add_table(doc, ["运维领域", "设计要求", "工作量建议"], [
        ["环境规划", "至少规划开发、测试、预生产、生产四类环境；生产采用应用服务与数据库分离部署，数据库主备。", "12-18 人日"],
        ["国产化适配", "适配达梦数据库、国产服务器操作系统、国产中间件或 Nginx/应用服务部署要求，提前验证驱动和 SQL 兼容性。", "20-35 人日"],
        ["配置管理", "数据库连接、文件存储、外部接口、导入目录、日志级别、任务开关全部配置化，避免写死在代码中。", "8-12 人日"],
        ["CI/CD", "建立构建、单元测试、打包、制品归档、部署脚本流程；生产部署需支持版本标记和回滚。", "15-25 人日"],
        ["日志审计", "区分业务操作日志、登录日志、接口日志、导入导出日志、系统错误日志；关键单据状态变化必须可追溯。", "15-25 人日"],
        ["监控告警", "监控应用存活、接口耗时、数据库连接池、慢 SQL、磁盘空间、导入任务、定时任务和异常错误率。", "15-25 人日"],
        ["备份恢复", "数据库全量/增量备份，导入文件和导出附件备份，定期恢复演练，形成恢复时间目标。", "12-20 人日"],
        ["上线回滚", "上线前冻结窗口、迁移脚本、冒烟检查、业务抽样、失败回滚、旧系统只读策略均需预案。", "15-25 人日"],
        ["运维交接", "提供部署手册、配置手册、巡检手册、常见问题、数据恢复手册和接口联调手册。", "10-18 人日"],
    ], [1.15, 4.2, 1.15], fill=LIGHT_BLUE)
    callout(doc, "部署工作量口径", "部署与运维建设建议单独预留 120-200 人日。如果甲方要求等保测评、密评配合、信创测评、双活容灾或第三方运维平台对接，应另行增加专项工作量和报价。")

    heading(doc, "七、详细功能项清单", 1)
    add_para(doc, "以下清单按需求文档附录逐项展开。正常人日为完整开发、自测、联调配合口径；协同后人日为采用 Vibe Coding 方式后的执行估算，不包含甲方等待、接口资料延迟和重大需求变更。")
    make_landscape(doc)
    heading(doc, "详细功能项与工作量估算", 1)
    detail_rows = []
    for r in rows:
        detail_rows.append([
            r["idx"],
            r["module"],
            r["page"],
            r["kind"],
            r["function"] or "按页面说明",
            r["dev"],
            f"{r['lo']}-{r['hi']}",
            f"{r['vibe_lo']}-{r['vibe_hi']}",
            r["detail"],
        ])
    add_table(
        doc,
        ["序号", "模块", "功能页面", "类型", "现有功能", "开发内容", "正常人日", "协同后", "难点/备注"],
        detail_rows,
        [0.35, 0.95, 1.35, 0.75, 1.25, 2.1, 0.58, 0.58, 1.6],
        fill=LIGHT_BLUE,
        font_size=7.2,
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    heading(doc, "八、风险与边界", 1)
    add_table(doc, ["风险点", "表现", "建议处理"], [
        ["需求边界不清", "文档中存在“待明确”“删除”“整合”项，容易反复变更。", "在开发前形成冻结版功能清单，变更走审批和工期调整。"],
        ["接口资料缺失", "数据湖、二维码、胆款、电子签章等接口材料未完全明确。", "一期只做接口框架和已确认接口，未确认部分单独列二期。"],
        ["报表口径争议", "同一报表集团/分公司/票务室口径不同，数字验收困难。", "先用旧系统样例数据签字确认，再开发。"],
        ["历史数据异常", "旧系统长期运行导致脏数据和状态不一致。", "迁移前做数据体检，异常数据由业务确认处理规则。"],
        ["AI 生成代码质量波动", "常规代码效率高，但核心业务可能遗漏隐性规则。", "核心交易类代码必须人工 review，并要求测试覆盖状态流转和金额计算。"],
    ], [1.35, 2.45, 2.7], fill=LIGHT)

    heading(doc, "九、结论", 1)
    add_para(doc, f"本项目完整实施建议按 10-12 人、8-10 个月规划；采用 Vibe Coding 协同后，可把常规代码、测试样例、文档和重复性页面工作压缩，推荐按 7-9 人、7-8 个月作为进取计划。部署与可运维性需同步纳入一期，额外预留 120-200 人日。即便采用协同开发，票据库存、票单结算、报表口径、历史数据迁移和上线割接仍是项目成败关键，需要架构师、业务负责人、测试负责人和运维负责人持续把关。")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
